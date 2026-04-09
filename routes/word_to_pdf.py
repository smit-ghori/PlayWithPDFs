import os
import shutil
import zipfile
import subprocess
import tempfile
import uuid
from io import BytesIO
from flask import Blueprint, render_template, request, send_file
from werkzeug.utils import secure_filename

from docx import Document
from docx.oxml.ns import qn
from PIL import Image as PILImage
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

word_to_pdf_bp = Blueprint("word_to_pdf", __name__)


def get_libreoffice_path():
    candidates = [
        os.environ.get("LIBREOFFICE_PATH"),
        shutil.which("soffice"),
        shutil.which("libreoffice"),
        "/usr/lib/libreoffice/program/soffice",
        "/usr/bin/soffice",
    ]

    for candidate in candidates:
        if candidate and os.path.isfile(candidate):
            return candidate

    raise RuntimeError(
        "LibreOffice not found. Install libreoffice and ensure 'soffice' is on PATH, or set LIBREOFFICE_PATH."
    )


def extract_images_from_run(run, temp_dir):
    images = []

    for blip in run._element.xpath('.//*[local-name() = "blip"]'):
        embed_id = blip.get(qn("r:embed"))
        if not embed_id:
            continue

        try:
            image_part = run.part.related_part(embed_id)
            image_bytes = image_part.blob
            format_name = PILImage.open(BytesIO(image_bytes)).format.lower()
            image_name = f"image_{uuid.uuid4().hex}.{format_name}"
            image_path = os.path.join(temp_dir, image_name)
            with open(image_path, "wb") as image_file:
                image_file.write(image_bytes)
            images.append(image_path)
        except Exception:
            continue

    return images


def paragraph_to_flowables(paragraph, styles, temp_dir):
    flowables = []
    text_buffer = ""
    alignment = TA_LEFT

    if paragraph.alignment:
        if paragraph.alignment == 1:
            alignment = TA_CENTER
        elif paragraph.alignment == 2:
            alignment = TA_RIGHT
        elif paragraph.alignment == 3:
            alignment = TA_JUSTIFY

    style = styles["Normal"]
    if paragraph.runs:
        first_run = paragraph.runs[0]
        if first_run.bold:
            style = ParagraphStyle(
                "BoldStyle",
                parent=styles["Normal"],
                fontSize=11,
                fontName="Helvetica-Bold",
                alignment=alignment,
            )
        elif first_run.italic:
            style = ParagraphStyle(
                "ItalicStyle",
                parent=styles["Normal"],
                fontSize=11,
                fontName="Helvetica-Oblique",
                alignment=alignment,
            )

    for run in paragraph.runs:
        image_paths = extract_images_from_run(run, temp_dir)
        if image_paths:
            if text_buffer.strip():
                flowables.append(Paragraph(text_buffer, style))
                text_buffer = ""

            for image_path in image_paths:
                try:
                    img = PILImage.open(image_path)
                    aspect = img.height / img.width if img.width else 1
                    max_width = 6 * inch
                    img_width = max_width
                    img_height = img_width * aspect
                    flowables.append(
                        Image(image_path, width=img_width, height=img_height)
                    )
                    flowables.append(Spacer(1, 0.2 * inch))
                except Exception:
                    continue

        if run.text:
            text_buffer += run.text

    if text_buffer.strip():
        flowables.append(Paragraph(text_buffer, style))
    elif not flowables:
        flowables.append(Spacer(1, 0.2 * inch))

    return flowables


def convert_word_to_pdf(input_file, output_dir=None):
    if not os.path.exists(input_file):
        raise FileNotFoundError(f"No such file: {input_file}")

    if output_dir is None:
        output_dir = os.path.dirname(input_file)

    pdf_filename = os.path.splitext(os.path.basename(input_file))[0] + ".pdf"
    pdf_path = os.path.join(output_dir, pdf_filename)

    temp_dir = tempfile.mkdtemp(prefix="docx_to_pdf_")
    try:
        doc = Document(input_file)
        story = []
        styles = getSampleStyleSheet()

        for paragraph in doc.paragraphs:
            story.extend(paragraph_to_flowables(paragraph, styles, temp_dir))

        for table in doc.tables:
            data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                data.append(row_data)

            if data:
                t = Table(data, colWidths=[1.5 * inch] * len(data[0]) if data else [])
                t.setStyle(
                    TableStyle(
                        [
                            ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
                            ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                            ("FONTSIZE", (0, 0), (-1, 0), 10),
                            ("BOTTOMPADDING", (0, 0), (-1, 0), 12),
                            ("BACKGROUND", (0, 1), (-1, -1), colors.beige),
                            ("GRID", (0, 0), (-1, -1), 1, colors.black),
                        ]
                    )
                )
                story.append(t)
                story.append(Spacer(1, 0.3 * inch))

        pdf_doc = SimpleDocTemplate(pdf_path, pagesize=letter)
        pdf_doc.build(story)
        return pdf_path
    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)


def convert_docx_to_pdf(input_file, output_dir=None):
    try:
        libreoffice_path = get_libreoffice_path()
        if output_dir is None:
            output_dir = os.path.dirname(input_file)

        subprocess.run(
            [
                libreoffice_path,
                "--headless",
                "--convert-to",
                "pdf",
                input_file,
                "--outdir",
                output_dir,
            ],
            check=True,
            capture_output=True,
            text=True,
        )

        pdf_filename = os.path.splitext(os.path.basename(input_file))[0] + ".pdf"
        return os.path.join(output_dir, pdf_filename)
    except RuntimeError as e:
        if "LibreOffice not found" in str(e):
            return convert_word_to_pdf(input_file, output_dir)
        raise
    except subprocess.CalledProcessError as e:
        raise RuntimeError(
            f"LibreOffice conversion failed: {e.stderr or e.stdout or str(e)}"
        ) from e
    except Exception as e:
        raise RuntimeError(f"Failed to convert {input_file} to PDF: {e}") from e


@word_to_pdf_bp.route("/word_to_pdf", methods=["GET", "POST"])
def word_to_pdf():
    if request.method == "POST":
        files = request.files.getlist("files")

        if not files:
            return "No files uploaded", 400

        allowed_extensions = {".docx"}
        for file in files:
            file_ext = os.path.splitext(file.filename)[1].lower()
            if file_ext not in allowed_extensions:
                return (
                    f"Invalid file type: {file.filename}. Only .docx files are supported.",
                    400,
                )

        converted_files = []
        with tempfile.TemporaryDirectory() as temp_dir:
            for file in files:
                filename = secure_filename(file.filename)
                if not filename:
                    return "Invalid file name", 400

                input_path = os.path.join(temp_dir, filename)
                file.save(input_path)

                pdf_path = convert_docx_to_pdf(input_path, temp_dir)
                converted_files.append((pdf_path, os.path.basename(pdf_path)))

            if len(converted_files) == 1:
                pdf_path, pdf_filename = converted_files[0]
                with open(pdf_path, "rb") as pdf_file:
                    pdf_stream = BytesIO(pdf_file.read())
                pdf_stream.seek(0)
                return send_file(
                    pdf_stream,
                    as_attachment=True,
                    download_name=pdf_filename,
                    mimetype="application/pdf",
                )

            zip_stream = BytesIO()
            with zipfile.ZipFile(zip_stream, "w", zipfile.ZIP_DEFLATED) as zipf:
                for pdf_path, pdf_filename in converted_files:
                    zipf.write(pdf_path, pdf_filename)

            zip_stream.seek(0)
            return send_file(
                zip_stream,
                as_attachment=True,
                download_name="converted_pdfs.zip",
                mimetype="application/zip",
            )

    return render_template("word_to_pdf.html")
